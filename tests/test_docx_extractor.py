import os
import tempfile
import docx
from src.extractors.docx_extractor import DocxExtractor

def test_docx_extraction():
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "sample.docx")
        doc = docx.Document()
        doc.add_heading("Project Proposal Template", level=1)
        doc.add_paragraph("This is a paragraph detailing project requirements.")

        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Phase"
        table.rows[0].cells[1].text = "Duration"
        table.rows[1].cells[0].text = "Phase 1"
        table.rows[1].cells[1].text = "2 Weeks"

        doc.save(docx_path)

        extractor = DocxExtractor()
        features = extractor.extract(docx_path)

        assert features.filename == "sample.docx"
        assert "Project Proposal Template" in features.headings
        assert len(features.tables) == 1
        assert features.tables[0].headers == ["Phase", "Duration"]
        assert "paragraph" in features.full_text.lower()
