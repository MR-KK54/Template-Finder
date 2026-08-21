import os
import pytest
from src.models import DocumentFeatures, PageData
from src.matcher.hybrid_comparator import HybridComparator
from src.matcher.sequence_matcher import (
    full_text_order_ratio,
    ordered_list_ratio,
    table_sequence_ratio,
    header_footer_ratio,
)
from src.models import TableData


def _make_doc(filepath: str, filename: str, full_text: str, headings: list,
              paragraphs: list = None, keywords: set = None, pages: int = 1,
              tables: list = None, headers_footers: list = None,
              file_hash: str = "h") -> DocumentFeatures:
    return DocumentFeatures(
        filepath=filepath,
        filename=filename,
        folder_name=os.path.basename(os.path.dirname(filepath)),
        file_size=5000,
        last_modified=1700000000.0,
        file_hash=file_hash,
        full_text=full_text,
        headings=headings,
        paragraphs=paragraphs or [],
        tables=tables or [],
        lists=[],
        headers_footers=headers_footers or [],
        keywords=set(keywords or []),
        page_count=pages,
        section_count=max(1, len(headings) + 1),
        pages=[PageData(page_num=1, text=full_text, headings=headings)]
    )


def test_sequence_matcher_ordered_vs_reversed():
    a = "Introduction to the study. Scope of work. Payment terms and conditions."
    b_same = "Introduction to the study. Scope of work. Payment terms and conditions."
    b_reversed = "Payment terms and conditions. Scope of work. Introduction to the study."

    same_score = full_text_order_ratio(a, b_same)
    reversed_score = full_text_order_ratio(a, b_reversed)
    assert same_score > reversed_score
    assert same_score > 80.0

    h1 = ["Introduction", "Scope of Work", "Payment Terms"]
    h2_reversed = ["Payment Terms", "Scope of Work", "Introduction"]
    ratio, matched = ordered_list_ratio(h1, h2_reversed)
    assert matched == 1
    assert ratio < 50.0


def test_stage3_rejects_reordered_document():
    comparator = HybridComparator()
    pdf = _make_doc(
        "pdf.pdf", "pdf.pdf",
        "Section A covers the introduction. Section B describes the scope. Section C defines payment terms.",
        ["Section A", "Section B", "Section C"],
        keywords=["introduction", "scope", "payment", "terms"],
        pages=1,
    )
    reordered = _make_doc(
        "reordered.docx", "reordered.docx",
        "Section C defines payment terms. Section B describes the scope. Section A covers the introduction.",
        ["Section C", "Section B", "Section A"],
        keywords=["introduction", "scope", "payment", "terms"],
        pages=1,
    )
    res = comparator.compare(pdf, reordered)
    assert res.rejected is True
    assert res.rejected_reason


def test_partial_export_from_large_doc_is_verified_source():
    """A 1-page PDF exported from a 20-page Word document must be the verified
    perfect source even though symmetric page similarity is very low."""
    comparator = HybridComparator()
    pdf = _make_doc(
        "pdf.pdf", "pdf.pdf",
        "Clinical study informed consent form for patient participation in a phase 2 trial.",
        ["Informed Consent"], keywords=["clinical", "consent", "patient"], pages=1,
    )
    large_source = _make_doc(
        "large_source.docx", "large_source.docx",
        "Clinical study informed consent form for patient participation in a phase 2 trial.",
        ["Informed Consent"], keywords=["clinical", "consent", "patient"], pages=20,
    )
    res = comparator.compare(pdf, large_source)
    assert not res.rejected
    assert res.verified_source is True
    assert res.overall_score == 100.0
    assert res.confidence_score == 100.0


def test_structure_penalty_reduces_confidence_when_not_verified():
    comparator = HybridComparator()
    pdf_text = "The protocol describes the patient eligibility criteria, the dosing schedule, and the monitoring requirements for the trial."
    pdf = _make_doc(
        "pdf.pdf", "pdf.pdf", pdf_text,
        ["Protocol"], keywords=["protocol", "eligibility", "schedule", "monitoring"], pages=1,
    )
    doc_text = ("This study protocol defines the participant inclusion and exclusion criteria, the medication "
                "administration schedule, and the required safety monitoring obligations for the clinical investigation.")
    doc1 = _make_doc(
        "doc1.docx", "doc1.docx", doc_text,
        ["Protocol"], keywords=["protocol", "eligibility", "schedule", "monitoring"], pages=1, file_hash="h1",
    )
    doc4 = _make_doc(
        "doc4.docx", "doc4.docx", doc_text,
        ["Protocol"], keywords=["protocol", "eligibility", "schedule", "monitoring"], pages=4, file_hash="h2",
    )

    res1 = comparator.compare(pdf, doc1)
    res4 = comparator.compare(pdf, doc4)

    assert not res1.rejected and not res4.rejected
    assert res1.verified_source is False and res4.verified_source is False
    assert res1.structure_score > res4.structure_score
    assert res1.confidence_score > res4.confidence_score


def test_stage_reports_are_always_present():
    comparator = HybridComparator()
    pdf = _make_doc("pdf.pdf", "pdf.pdf", "Annual report content here.", ["Annual Report"])
    word = _make_doc("word.docx", "word.docx", "Annual report content here.", ["Annual Report"])
    res = comparator.compare(pdf, word)

    stages = [s.stage for s in res.stage_reports]
    assert any("Stage 3" in s for s in stages)
    assert any("Stage 4" in s for s in stages)
    assert any("Stage 5" in s for s in stages)
    assert any("Stage 6" in s for s in stages)
    assert len(res.debug_stages) == len(res.stage_reports)
    assert res.selection_reason


def test_compare_batch_accepted_first_rejected_flagged():
    comparator = HybridComparator()
    pdf = _make_doc(
        "pdf.pdf", "pdf.pdf",
        "Clinical study informed consent form for patient participation in a phase 2 trial.",
        ["Informed Consent"], keywords=["clinical", "consent", "patient"],
    )
    source = _make_doc(
        "source.docx", "source.docx",
        "Clinical study informed consent form for patient participation in a phase 2 trial.",
        ["Informed Consent"], keywords=["clinical", "consent", "patient"], file_hash="h1",
    )
    unrelated = _make_doc(
        "unrelated.docx", "unrelated.docx",
        "Shipping invoice and delivery terms for medical supplies and equipment.",
        ["Shipping Invoice"], keywords=["shipping", "invoice", "delivery"], file_hash="h2",
    )

    results = comparator.compare_batch(pdf, [unrelated, source], top_k=2)
    assert len(results) == 2
    assert results[0].rejected is False
    assert results[0].word_file_name == "source.docx"
    assert results[1].rejected is True
    assert results[1].rejected_reason


def test_table_sequence_ratio_order_sensitive():
    t1 = [TableData(rows=2, cols=2, headers=["A", "B"], flat_text="A B"), TableData(rows=2, cols=2, headers=["C", "D"], flat_text="C D")]
    t2_same = [TableData(rows=2, cols=2, headers=["A", "B"], flat_text="A B"), TableData(rows=2, cols=2, headers=["C", "D"], flat_text="C D")]
    t2_reversed = [TableData(rows=2, cols=2, headers=["C", "D"], flat_text="C D"), TableData(rows=2, cols=2, headers=["A", "B"], flat_text="A B")]
    same_score = table_sequence_ratio(t1, t2_same)
    reversed_score = table_sequence_ratio(t1, t2_reversed)
    assert same_score > 90.0
    assert same_score > reversed_score
    assert reversed_score < 80.0


def test_header_footer_ratio():
    assert header_footer_ratio([], []) == 100.0
    assert header_footer_ratio(["Confidential"], ["Confidential"]) == 100.0
    assert header_footer_ratio(["Confidential"], ["Public"]) == 0.0


def test_front_page_match_verifies_when_later_pages_differ():
    """The PDF's first page is fully contained in the Word document but later
    pages differ -> verified as the original source via the front page."""
    comparator = HybridComparator()
    front_text = ("3K JEEVA BARATHI NAGAR BALAJI AVENUE ASHOK NAGAR COIMBATORE PHONE 9790118995. "
                  "To ICEGATE COIMBATORE AIRPORT. Dear Sir, SUB: APPLICATION FOR GRANT OF RODTEP LICENSE. "
                  "Kindly inform you that we are sending the following documents for availing benefit.")
    later_text = ("This is entirely different content that does not exist in the target Word document. "
                  "It contains unrelated regulatory statements and additional annexures that were added "
                  "after the export was generated from the original template document.")
    pdf = DocumentFeatures(
        filepath="letter.pdf", filename="letter.pdf", folder_name=".",
        file_size=1000, last_modified=1700000000.0, file_hash="hp",
        full_text=front_text + "\n\n" + later_text,
        headings=["SUB: APPLICATION FOR GRANT OF RODTEP LICENSE"],
        paragraphs=[front_text, later_text],
        keywords=set(w for w in front_text.split() if len(w) >= 4),
        page_count=2,
        section_count=2,
        pages=[
            PageData(page_num=1, text=front_text, headings=["SUB: APPLICATION FOR GRANT OF RODTEP LICENSE"]),
            PageData(page_num=2, text=later_text, headings=[]),
        ],
    )
    word = _make_doc(
        "source.docx", "source.docx", front_text,
        ["SUB: APPLICATION FOR GRANT OF RODTEP LICENSE"],
        paragraphs=[front_text], keywords=["rodtep", "icegate", "licence"], pages=1, file_hash="hw",
    )

    res = comparator.compare(pdf, word)
    assert not res.rejected
    assert res.verified_source is True
    assert res.match_basis == "front page"
    assert res.front_coverage >= 85.0
    assert res.text_coverage < 85.0
    assert res.overall_score == 100.0


def test_front_page_blank_falls_back_to_full_document():
    """A blank/short front page must not block full-document verification."""
    comparator = HybridComparator()
    body = ("Clinical study informed consent form for patient participation in a phase 2 trial conducted "
            "at the research hospital. The participant will receive the investigational product for twelve weeks.")
    pdf = DocumentFeatures(
        filepath="consent.pdf", filename="consent.pdf", folder_name=".",
        file_size=1000, last_modified=1700000000.0, file_hash="hp",
        full_text=body, headings=["Informed Consent"],
        paragraphs=[body], keywords={"clinical", "consent", "patient"},
        page_count=1, section_count=1,
        pages=[PageData(page_num=1, text="", headings=[])],
    )
    word = _make_doc(
        "consent.docx", "consent.docx", body,
        ["Informed Consent"], paragraphs=[body], keywords=["clinical", "consent", "patient"], pages=1, file_hash="hw",
    )

    res = comparator.compare(pdf, word)
    assert not res.rejected
    assert res.verified_source is True
    assert res.match_basis in ("full document", "both")
    assert res.overall_score == 100.0


def test_docx_extractor_includes_header_text_in_full_text():
    import tempfile
    from src.extractors.docx_extractor import DocxExtractor
    import docx as docx_lib

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "letterhead.docx")
        d = docx_lib.Document()
        d.add_heading("SUB: APPLICATION FOR GRANT OF RODTEP LICENSE", level=1)
        d.add_paragraph("Kindly inform you that we are sending the following documents for availing benefit.")
        header = d.sections[0].header
        header.paragraphs[0].text = "3K JEEVA BARATHI NAGAR BALAJI AVENUE ASHOK NAGAR COIMBATORE PHONE 9790118995"
        d.save(docx_path)

        features = DocxExtractor().extract(docx_path)
        assert any("JEEVA" in hf for hf in features.headers_footers)
        assert "JEEVA" in features.full_text.upper()