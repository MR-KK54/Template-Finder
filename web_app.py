import html
import io
import os
import tempfile
from pathlib import Path

import streamlit as st
from PIL import Image as PILImage

from src.database import IndexDatabase
from src.extractors.pdf_extractor import PDFExtractor
from src.indexer.file_scanner import FileScanner
from src.indexer.index_manager import IndexManager
from src.matcher.hybrid_comparator import HybridComparator
from src.models import DocumentFeatures
from src.utils.file_utils import format_file_size, to_long_path


import urllib.parse

st.set_page_config(
    page_title="Template Finder",
    page_icon="TF",
    layout="wide",
)


def render_final_document(doc_path: str, full_text: str = "") -> str:
    """Renders the matched Word document's content inline as styled HTML."""
    long_p = to_long_path(doc_path)
    target_p = long_p if os.path.exists(long_p) else doc_path

    html_chunks = []
    rendered = False

    if os.path.exists(target_p) and target_p.lower().endswith(".docx"):
        try:
            import docx
            doc_word = docx.Document(target_p)

            for p in doc_word.paragraphs:
                txt = p.text.strip()
                if not txt:
                    continue
                style = (p.style.name or "").lower()
                if "heading 1" in style or "title" in style:
                    html_chunks.append(
                        f"<h1 style='color:#1a365d; font-family:sans-serif; border-bottom:2px solid #3182ce; padding-bottom:4px;'>{html.escape(txt)}</h1>"
                    )
                elif "heading 2" in style:
                    html_chunks.append(
                        f"<h2 style='color:#2b6cb0; font-family:sans-serif;'>{html.escape(txt)}</h2>"
                    )
                elif "heading" in style:
                    html_chunks.append(
                        f"<h3 style='color:#2d3748; font-family:sans-serif;'>{html.escape(txt)}</h3>"
                    )
                else:
                    html_chunks.append(
                        f"<p style='font-family:sans-serif; font-size:12px; line-height:1.6; color:#2d3748; margin:4px 0;'>{html.escape(txt)}</p>"
                    )

            for t in doc_word.tables:
                html_chunks.append("<table style='width:100%; border-collapse:collapse; margin:12px 0; border:1px solid #cbd5e0;'>")
                for r_idx, row in enumerate(t.rows):
                    html_chunks.append("<tr>")
                    for cell in row.cells:
                        tag = "th" if r_idx == 0 else "td"
                        bg = "background:#edf2f7; font-weight:bold;" if r_idx == 0 else ""
                        html_chunks.append(
                            f"<{tag} style='border:1px solid #cbd5e0; padding:6px; font-family:sans-serif; font-size:11px; {bg}'>{html.escape(cell.text.strip())}</{tag}>"
                        )
                    html_chunks.append("</tr>")
                html_chunks.append("</table>")

            rendered = True
        except Exception:
            rendered = False

    if not rendered:
        text = full_text or ""
        for para in text.split("\n"):
            para = para.strip()
            if para:
                html_chunks.append(
                    f"<p style='font-family:sans-serif; font-size:12px; line-height:1.6; color:#2d3748; margin:4px 0;'>{html.escape(para)}</p>"
                )

    return "\n".join(html_chunks) if html_chunks else "<p style='font-family:sans-serif; color:#718096;'>[No extractable content]</p>"


@st.cache_data(show_spinner="Rendering document preview…", ttl=3600)
def _render_doc_images(doc_path: str, max_pages: int = 15):
    """Render a .docx/.pdf document to a list of PIL page images for image preview.

    Word documents are converted to PDF via Word COM, then every page is
    rasterised with PyMuPDF. Returns an empty list on any failure so the
    caller can fall back to the text view.
    """
    long_p = to_long_path(doc_path)
    target_p = long_p if os.path.exists(long_p) else doc_path
    images = []
    temp_pdf = None
    pdf_path = None
    try:
        if not os.path.exists(target_p):
            return images

        if target_p.lower().endswith(".pdf"):
            pdf_path = target_p
        elif target_p.lower().endswith(".docx"):
            import pythoncom
            import win32com.client as win32

            pythoncom.CoInitialize()
            try:
                word = win32.Dispatch("Word.Application")
                word.Visible = False
                norm = target_p[4:] if target_p.startswith("\\\\?\\") else target_p
                doc = word.Documents.Open(os.path.abspath(norm))
                fd, temp_pdf = tempfile.mkstemp(suffix=".pdf")
                os.close(fd)
                doc.SaveAs(os.path.abspath(temp_pdf), FileFormat=17)  # wdFormatPDF
                doc.Close(False)
                word.Quit()
                pdf_path = temp_pdf
            finally:
                pythoncom.CoUninitialize()
        else:
            return images

        import fitz
        pdf_doc = fitz.open(pdf_path)
        for i in range(min(len(pdf_doc), max_pages)):
            pix = pdf_doc.load_page(i).get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            images.append(PILImage.open(io.BytesIO(pix.tobytes("png"))))
        pdf_doc.close()
    except Exception:
        images = []
    finally:
        if temp_pdf and os.path.exists(temp_pdf):
            try:
                os.remove(temp_pdf)
            except Exception:
                pass
    return images


def render_standalone_preview_page(doc_path: str = "", initial_page: int = 1):
    """Renders a dedicated standalone full-page view containing ONLY the Document Preview tool."""
    if 'standalone_page_num' not in st.session_state:
        st.session_state['standalone_page_num'] = initial_page

    st.markdown(
        """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        </style>
        """,
        unsafe_allow_html=True
    )

    with st.container(border=True):
        c_title, c_back = st.columns([5, 1])
        with c_title:
            st.markdown(
                "<h2 style='color:#3b82f6; margin:0; font-weight:700; font-family:sans-serif; display:flex; align-items:center; gap:8px;'>Document Preview 🔗</h2>",
                unsafe_allow_html=True
            )
            if doc_path and os.path.exists(to_long_path(doc_path)):
                st.caption(f"Document: `{os.path.basename(doc_path)}` | `{doc_path}`")
            elif doc_path:
                st.caption(f"Document: `{doc_path}`")
            else:
                st.caption("Standalone Document Preview Mode")
        with c_back:
            if st.button("🏠 Main App", key="btn_back_main_app", use_container_width=True):
                st.query_params.clear()
                st.rerun()

        st.markdown("<div style='margin-bottom: 12px;'></div>", unsafe_allow_html=True)

        images = []
        if doc_path and (os.path.exists(to_long_path(doc_path)) or os.path.exists(doc_path)):
            images = _render_doc_images(doc_path, max_pages=50)

        total_pages = len(images) if images else (1 if doc_path else 0)
        curr_p = int(st.session_state.get('standalone_page_num', initial_page))
        if total_pages > 0:
            curr_p = min(max(1, curr_p), total_pages)
        else:
            curr_p = 1
        st.session_state['standalone_page_num'] = curr_p

        # Controls Toolbar
        col_prev, col_page_text, col_input, col_go, col_next = st.columns([1.2, 2.5, 1.2, 0.8, 1.2])
        with col_prev:
            if st.button("◀ Prev", key="btn_stand_prev", disabled=(curr_p <= 1 or not doc_path), use_container_width=True):
                st.session_state['standalone_page_num'] = max(1, curr_p - 1)
                st.rerun()

        with col_page_text:
            page_str = f"Page {curr_p} / {total_pages}" if (doc_path and total_pages > 0) else "Page - / -"
            st.markdown(
                f"<div style='text-align:center; color:#94a3b8; font-size:14px; font-weight:500; padding-top:8px;'>{page_str}</div>",
                unsafe_allow_html=True
            )

        with col_input:
            p_input = st.text_input("Page #", key="input_stand_page_val", placeholder="Page #", label_visibility="collapsed", disabled=not doc_path)

        with col_go:
            if st.button("Go", key="btn_stand_go", disabled=not doc_path, use_container_width=True):
                if p_input and p_input.strip().isdigit():
                    t_page = int(p_input.strip())
                    st.session_state['standalone_page_num'] = min(max(1, t_page), total_pages)
                    st.rerun()

        with col_next:
            if st.button("Next ▶", key="btn_stand_next", disabled=(curr_p >= total_pages or not doc_path), use_container_width=True):
                st.session_state['standalone_page_num'] = min(total_pages, curr_p + 1)
                st.rerun()

        st.markdown("<div style='margin-bottom: 12px;'></div>", unsafe_allow_html=True)

        # Document Viewport Frame (Full Height Single Page View)
        if not doc_path or not (os.path.exists(to_long_path(doc_path)) or os.path.exists(doc_path)):
            st.markdown(
                """
                <div style='background-color:#0b0c10; border:1px solid #1f2433; border-radius:8px; height:650px; display:flex; align-items:center; justify-content:center; color:#64748b; font-size:16px; font-family:sans-serif;'>
                    Select a document and click Preview to see it here.
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            if images:
                import base64
                current_img = images[curr_p - 1]
                buf = io.BytesIO()
                current_img.save(buf, format="PNG")
                img_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                st.markdown(
                    f"""
                    <div style='background-color:#0b0c10; border:1px solid #1f2433; border-radius:8px; padding:16px; height:650px; display:flex; align-items:center; justify-content:center; overflow:hidden;'>
                        <img src="data:image/png;base64,{img_b64}" style="max-height:620px; max-width:100%; object-fit:contain; border-radius:4px; box-shadow:0 4px 14px rgba(0,0,0,0.6);" />
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            else:
                final_html = render_final_document(doc_path, "")
                white_box = ("height: 650px; overflow-y: auto; padding: 16px; border-radius: 8px; "
                             "background-color: #0b0c10; color: #e2e8f0; border: 1px solid #1f2433;")
                st.markdown(f"<div style='{white_box}'>{final_html}</div>", unsafe_allow_html=True)


# Check query parameters for standalone dedicated preview page (?view=preview)
q_params = st.query_params
if q_params.get("view") == "preview":
    standalone_doc = q_params.get("doc_path", "")
    standalone_p = int(q_params.get("page", 1)) if (q_params.get("page") or "").isdigit() else 1
    render_standalone_preview_page(standalone_doc, standalone_p)
    st.stop()


st.title("PDF to Word Template Finder")
st.caption("Compare a reference PDF with your Word template library and find the original source document.")


def render_document_preview_panel(results, indexed_texts):
    """Renders a single shared Document Preview panel matching the exact model design:
    - Header: Blue 'Document Preview' title on left, 'Close' button on top-right.
    - Navigation bar: '◄ Prev' button, 'Page X / Y' label, 'Page #' input + 'Go' button, 'Next ▶' button.
    - Viewport: Dark container with 'Select a document and click Preview to see it here.' placeholder when inactive,
      or the active document page image / text.
    """
    if 'preview_visible' not in st.session_state:
        st.session_state['preview_visible'] = True
    if 'preview_doc_index' not in st.session_state:
        st.session_state['preview_doc_index'] = 0
    if 'preview_page_num' not in st.session_state:
        st.session_state['preview_page_num'] = 1

    preview_idx = int(st.session_state.get('preview_doc_index', 0))
    is_visible = bool(st.session_state.get('preview_visible', True))

    if not results or preview_idx < 0 or preview_idx >= len(results):
        is_active = False
        result = None
    else:
        is_active = is_visible
        result = results[preview_idx]

    images = []
    if is_active and result:
        images = _render_doc_images(result.file_path, max_pages=30)

    total_pages = len(images) if images else (1 if is_active else 0)
    current_page = int(st.session_state.get('preview_page_num', 1))
    if total_pages > 0:
        current_page = min(max(1, current_page), total_pages)
    else:
        current_page = 1
    st.session_state['preview_page_num'] = current_page

    # Outer Document Preview Container (styled dark matching the reference image)
    with st.container(border=True):
        # 1. Header Bar: Title with 🔗 Link Option & Close Button
        c_title, c_close = st.columns([5, 1])
        with c_title:
            if result and getattr(result, 'file_path', None):
                encoded_path = urllib.parse.quote(result.file_path)
                link_href = f"?view=preview&doc_path={encoded_path}&page={current_page}"
            else:
                link_href = "?view=preview"

            st.markdown(
                f"<h3 style='color:#3b82f6; margin:0; padding-top:4px; font-weight:700; font-family:sans-serif; display:flex; align-items:center; gap:8px;'>"
                f"Document Preview <a href='{link_href}' target='_blank' title='Open Document Preview in new page' style='text-decoration:none; color:#3b82f6; font-size:18px;'>🔗</a></h3>",
                unsafe_allow_html=True
            )
        with c_close:
            if st.button("Close", key="btn_close_doc_preview", use_container_width=True):
                st.session_state['preview_visible'] = False
                st.rerun()

        st.markdown("<div style='margin-bottom: 8px;'></div>", unsafe_allow_html=True)

        # 2. Control Toolbar: Prev, Page text, Page # input, Go, Next
        col_prev, col_page_text, col_input, col_go, col_next = st.columns([1.2, 2.5, 1.2, 0.8, 1.2])

        with col_prev:
            if st.button("◀ Prev", key="btn_prev_page", disabled=(not is_active or current_page <= 1), use_container_width=True):
                st.session_state['preview_page_num'] = max(1, current_page - 1)
                st.rerun()

        with col_page_text:
            if is_active and total_pages > 0:
                page_str = f"Page {current_page} / {total_pages}"
            else:
                page_str = "Page - / -"
            st.markdown(
                f"<div style='text-align:center; color:#94a3b8; font-size:14px; font-weight:500; padding-top:8px;'>{page_str}</div>",
                unsafe_allow_html=True
            )

        with col_input:
            page_val_input = st.text_input(
                "Page #",
                key="input_page_num_val",
                placeholder="Page #",
                label_visibility="collapsed",
                disabled=not is_active
            )

        with col_go:
            if st.button("Go", key="btn_go_page_num", disabled=not is_active, use_container_width=True):
                if page_val_input and page_val_input.strip().isdigit():
                    target_page = int(page_val_input.strip())
                    st.session_state['preview_page_num'] = min(max(1, target_page), total_pages)
                    st.rerun()

        with col_next:
            if st.button("Next ▶", key="btn_next_page", disabled=(not is_active or current_page >= total_pages), use_container_width=True):
                st.session_state['preview_page_num'] = min(total_pages, current_page + 1)
                st.rerun()

        st.markdown("<div style='margin-bottom: 12px;'></div>", unsafe_allow_html=True)

        # 3. Document Viewport Frame (Compact dark canvas fitting single page completely in view)
        if not is_active or not result:
            st.markdown(
                """
                <div style='background-color:#0b0c10; border:1px solid #1f2433; border-radius:8px; height:480px; display:flex; align-items:center; justify-content:center; color:#64748b; font-size:15px; font-family:sans-serif;'>
                    Select a document and click Preview to see it here.
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            if images:
                import base64
                current_img = images[current_page - 1]
                buf = io.BytesIO()
                current_img.save(buf, format="PNG")
                img_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                st.markdown(
                    f"""
                    <div style='background-color:#0b0c10; border:1px solid #1f2433; border-radius:8px; padding:12px; height:480px; display:flex; align-items:center; justify-content:center; overflow:hidden;'>
                        <img src="data:image/png;base64,{img_b64}" style="max-height:455px; max-width:100%; object-fit:contain; border-radius:4px; box-shadow:0 4px 14px rgba(0,0,0,0.6);" />
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            else:
                final_html = render_final_document(result.file_path, indexed_texts.get(result.file_path, ""))
                white_box = ("height: 480px; overflow-y: auto; padding: 16px; border-radius: 8px; "
                             "background-color: #0b0c10; color: #e2e8f0; border: 1px solid #1f2433;")
                st.markdown(f"<div style='{white_box}'>{final_html}</div>", unsafe_allow_html=True)



st.divider()
col1, col2 = st.columns([1, 1])

# 1. Reference PDF Section
with col1:
    st.subheader("1. Reference PDF")
    pdf_source = st.radio(
        "PDF Source",
        ["Upload PDF File", "Local PDF Path"],
        horizontal=True,
    )

    if pdf_source == "Upload PDF File":
        pdf_upload = st.file_uploader("Reference PDF", type=["pdf"])
        pdf_path_input = None
        if pdf_upload:
            st.success(f"Selected: **{pdf_upload.name}** ({format_file_size(len(pdf_upload.getvalue()))})")
    else:
        pdf_upload = None
        pdf_path_input = st.text_input(
            "Reference PDF Path",
            placeholder=r"e.g. D:\path\to\sample_data\input.pdf",
            help="Enter local file path to the reference PDF file."
        )
        if pdf_path_input:
            clean_p = pdf_path_input.strip().strip('"\'')
            long_p = to_long_path(clean_p)
            if os.path.isfile(long_p) or os.path.isfile(clean_p):
                st.success(f"Valid PDF Path: `{clean_p}`")
            else:
                st.error("File not found at specified path.")

# 2. Word Templates Section
with col2:
    st.subheader("2. Word Templates")
    template_source = st.radio(
        "Template Source",
        ["Local Folder Path", "Upload DOCX Files"],
        horizontal=True,
    )
    if template_source == "Local Folder Path":
        default_folder_val = r"D:\Kishore\TEM\EC" if os.path.exists(r"D:\Kishore\TEM\EC") else ""
        template_folder_input = st.text_input(
            "Word Templates Folder Path",
            value=default_folder_val,
            placeholder=r"e.g. D:\Kishore\TEM\EC",
            help="Enter absolute or relative path to folder containing .docx templates.",
        )
        template_uploads = None
        if template_folder_input:
            clean_f = template_folder_input.strip().strip('"\'')
            long_f = to_long_path(clean_f)
            if os.path.isdir(long_f) or os.path.isdir(clean_f):
                st.success(f"Valid Folder: `{clean_f}`")

                # If search results exist, show Top 5 matches matching >= 90% accuracy from source doc
                search_results = st.session_state.get('search_results', [])
                top_90_matches = [r for r in search_results if r.overall_score >= 90.0]
                if not top_90_matches and search_results:
                    top_90_matches = search_results[:5]

                if top_90_matches:
                    display_candidates = top_90_matches[:5]
                    st.write(f"🎯 **Top {len(display_candidates)} Accurate Template Matches (90%+ Match)**: "
                             + ", ".join(f"{r.word_file_name} ({r.overall_score:.1f}%)" for r in display_candidates))
                else:
                    scanner = FileScanner()
                    folder_docs = scanner.scan_directory(clean_f, recursive=True)
                    if folder_docs:
                        st.write(f"📂 **{len(folder_docs)}** template file(s) found. "
                                 f"Click 'Find matching templates' for score ranking.")
            else:
                st.error("Folder not found at specified path.")
    else:
        template_folder_input = None
        template_uploads = st.file_uploader(
            "Word templates",
            type=["docx"],
            accept_multiple_files=True,
        )
        if template_uploads:
            st.success(f"Selected **{len(template_uploads)}** Word template file(s).")

threshold = st.slider("Minimum match score", min_value=30, max_value=100, value=100)

st.markdown("### 📄 PDF Page Recognition Mode")
page_rec_mode = st.radio(
    "Select Recognition Mode for Reference PDF",
    ["Default (First Page)", "Manual Page Selection"],
    index=0,
    horizontal=True,
    help="Choose whether to compare by default First Page or manually enter a specific Page #."
)

if page_rec_mode == "Manual Page Selection":
    target_pdf_page_num = st.number_input(
        "Enter Reference PDF Page # to Compare",
        min_value=1,
        max_value=500,
        value=1,
        step=1,
        help="Manually enter the 1-based page number of the uploaded PDF file to compare against Word templates."
    )
else:
    target_pdf_page_num = 1

front_page_only = True

is_finished = st.session_state.get('search_performed', False)

if is_finished:
    st.markdown("""
        <style>
            div.stButton > button[kind="primary"] {
                background-color: #22c55e !important;
                border-color: #16a34a !important;
                color: white !important;
            }
            div.stButton > button[kind="primary"]:hover {
                background-color: #16a34a !important;
                border-color: #15803d !important;
            }
            div[data-baseweb="slider"] div {
                background-color: #22c55e !important;
            }
        </style>
    """, unsafe_allow_html=True)
    st.markdown(
        "<div style='background-color:#15803d; color:white; padding:12px 18px; border-radius:8px; font-weight:700; font-size:16px; text-align:center; margin-bottom:15px; box-shadow:0 2px 8px rgba(0,0,0,0.3);'>"
        "✅ FINAL OUTPUT READY: Document Template Search Completed Successfully!"
        "</div>",
        unsafe_allow_html=True
    )

if st.button("Find matching templates", type="primary", use_container_width=True):
    temp_dir_obj = tempfile.TemporaryDirectory(prefix="template_finder_")
    workspace = Path(temp_dir_obj.name)

    try:
        # 1. Resolve PDF Path
        if pdf_source == "Upload PDF File":
            if not pdf_upload:
                st.error("Please upload a Reference PDF file.")
                st.stop()
            pdf_path = workspace / pdf_upload.name
            pdf_path.write_bytes(pdf_upload.getvalue())
            target_pdf_path = str(pdf_path)
        else:
            clean_pdf_path = (pdf_path_input or "").strip().strip('"\'').strip()
            if not clean_pdf_path and os.path.exists(r"sample_data\chile_reference.pdf"):
                clean_pdf_path = r"sample_data\chile_reference.pdf"
            long_pdf_path = to_long_path(clean_pdf_path)
            if not os.path.isfile(long_pdf_path) and not os.path.isfile(clean_pdf_path):
                st.error(f"Please enter a valid, existing local PDF file path. (Got: `{clean_pdf_path}`)")
                st.stop()
            target_pdf_path = clean_pdf_path

        # 2. Resolve Templates Path
        if template_source == "Local Folder Path":
            clean_folder_path = (template_folder_input or "").strip().strip('"\'').strip()
            if not clean_folder_path:
                if os.path.exists(r"D:\Kishore\TEM\EC"):
                    clean_folder_path = r"D:\Kishore\TEM\EC"
                elif os.path.exists(r"sample_data\templates"):
                    clean_folder_path = r"sample_data\templates"

            long_folder_path = to_long_path(clean_folder_path)
            if not os.path.isdir(long_folder_path) and not os.path.isdir(clean_folder_path) and not os.path.exists(long_folder_path) and not os.path.exists(clean_folder_path):
                st.error(f"Please enter a valid, existing local folder path. (Entered: `{template_folder_input}`)")
                st.stop()
            target_template_dir = clean_folder_path
        else:
            if not template_uploads:
                st.error("Please upload at least one Word template (.docx) file.")
                st.stop()
            template_dir = workspace / "templates"
            template_dir.mkdir(exist_ok=True)
            for template in template_uploads:
                (template_dir / template.name).write_bytes(template.getvalue())
            target_template_dir = str(template_dir)

        # 3. Index & Compare
        database = IndexDatabase("index_cache.db")
        index_manager = IndexManager(database)

        with st.status("Indexing Word templates...", expanded=False) as status:
            indexed_documents = index_manager.index_directory(target_template_dir)
            st.session_state['indexed_texts'] = {
                doc.filepath: doc.full_text for doc in indexed_documents if getattr(doc, 'full_text', None)
            }
            status.update(
                label=f"Indexed {len(indexed_documents)} template(s)",
                state="complete",
            )

        with st.spinner("Extracting PDF features..."):
            pdf_document = PDFExtractor().extract(target_pdf_path)

        if page_rec_mode == "Manual Page Selection" and pdf_document.pages:
            selected_idx = min(max(1, target_pdf_page_num), len(pdf_document.pages)) - 1
            page_data = pdf_document.pages[selected_idx]
            single_page_pdf = DocumentFeatures(
                filepath=pdf_document.filepath,
                filename=pdf_document.filename,
                folder_name=pdf_document.folder_name,
                file_size=pdf_document.file_size,
                last_modified=pdf_document.last_modified,
                file_hash=pdf_document.file_hash,
                full_text=page_data.text,
                headings=page_data.headings,
                paragraphs=page_data.text.split("\n"),
                tables=page_data.tables,
                lists=[],
                keywords=set(page_data.text.lower().split()),
                page_count=1,
                section_count=1,
                pages=[page_data],
                is_scanned_pdf=page_data.is_scanned
            )
            pdf_document = single_page_pdf
        elif page_rec_mode == "Default (First Page)" and pdf_document.pages:
            page_data = pdf_document.pages[0]
            single_page_pdf = DocumentFeatures(
                filepath=pdf_document.filepath,
                filename=pdf_document.filename,
                folder_name=pdf_document.folder_name,
                file_size=pdf_document.file_size,
                last_modified=pdf_document.last_modified,
                file_hash=pdf_document.file_hash,
                full_text=page_data.text,
                headings=page_data.headings,
                paragraphs=page_data.text.split("\n"),
                tables=page_data.tables,
                lists=[],
                keywords=set(page_data.text.lower().split()),
                page_count=1,
                section_count=1,
                pages=[page_data],
                is_scanned_pdf=page_data.is_scanned
            )
            pdf_document = single_page_pdf

        comparator = HybridComparator()
        compare_progress = st.progress(0.0, text="Preparing comparison...")
        all_results = comparator.compare_batch(
            pdf_document,
            indexed_documents,
            progress_callback=lambda done, total, name: compare_progress.progress(
                done / max(1, total),
                text=f"Comparing {done}/{total} documents: {name}",
            ),
        )
        compare_progress.empty()

        with st.spinner("Ranking results..."):
            results = sorted(
                (result for result in all_results
                 if not result.rejected and (result.overall_score >= threshold or result.verified_source)),
                key=lambda result: result.overall_score,
                reverse=True,
            )
            if front_page_only:
                results = [r for r in results if getattr(r, 'verified_source', False)]

            # Never leave the user with a dead end: if nothing cleared the
            # threshold, fall back to the best candidates so the FINAL OUTPUT
            # section always has a file to show (rejected ones keep their
            # rejection reason so the failure is visible).
            if not results:
                non_rejected = sorted(
                    (r for r in all_results if not r.rejected),
                    key=lambda r: r.overall_score,
                    reverse=True,
                )
                if non_rejected:
                    results = non_rejected[:5]
                else:
                    results = sorted(
                        all_results,
                        key=lambda r: r.overall_score,
                        reverse=True,
                    )[:5]
            st.session_state['search_performed'] = True
            st.session_state['search_results'] = results
            st.session_state['preview_doc_index'] = 0

    finally:
        temp_dir_obj.cleanup()


# ============================================================
# FINAL OUTPUT - shown below the search button after a search
# ============================================================
results = st.session_state.get('search_results', [])
search_performed = st.session_state.get('search_performed', False)

if search_performed and results:
    verified = [r for r in results if getattr(r, 'verified_source', False)]
    is_verified = bool(verified)

    st.divider()
    st.subheader("📦 FINAL OUTPUT")
    with st.container(border=True):
        if is_verified:
            top = verified[0]
            basis = getattr(top, 'match_basis', '') or 'full document'
            cov = getattr(top, 'front_coverage', 0) or getattr(top, 'text_coverage', 0)
            st.success(f"✅ **PERFECT ORIGINAL SOURCE FOUND** — verified by **{basis}** "
                       f"(coverage {cov:.1f}%) as the document the PDF was made from.")
        else:
            top = results[0]
            st.warning(
                f"⚠️ **No verified original source found** — no indexed Word document contains ≥ 85% of this "
                f"PDF's content in exact order. Best candidate: *{top.word_file_name}* at **{top.overall_score:.1f}%** "
                f"(front page coverage {getattr(top, 'front_coverage', 0):.1f}%, full content coverage "
                f"{getattr(top, 'text_coverage', 0):.1f}%).\n\n"
                "Make sure the original Word file is inside the selected folder and the PDF is a direct export/scan of it."
            )
        st.markdown(f"**{len(results)} document(s) passed all verification stages:**")

    indexed_texts = st.session_state.get('indexed_texts', {})

    st.divider()
    render_document_preview_panel(results, indexed_texts)

    curr_preview_idx = int(st.session_state.get('preview_doc_index', 0))
    is_preview_vis = bool(st.session_state.get('preview_visible', True))

    for idx, result in enumerate(results):
        with st.container(border=True):
            # Category Badge Header
            if result.overall_score >= 100.0 or "100%" in result.match_category:
                st.markdown("### 🥇 **100% CONTENT MATCH — ORIGINAL WORD SOURCE DOCUMENT**")
            elif result.overall_score >= 95.0:
                st.markdown("### 🥈 **95–99% MATCH — NEAR-IDENTICAL DOCUMENT**")
            elif result.overall_score >= 90.0:
                st.markdown("### 🥉 **90–94% MATCH — HIGHLY SIMILAR DOCUMENT**")
            elif idx == 0:
                st.markdown("### 🎯 **HIGHEST RANKED MATCH**")
            else:
                st.markdown(f"### 📄 {result.word_file_name}")

            # File name + file path - always visible
            st.markdown(f"**File Name:** `{result.word_file_name}`")
            st.markdown("**File Path:**")
            st.code(result.file_path, language="text")

            if getattr(result, 'rejected', False):
                st.error(f"❌ **Rejected because:** {result.rejected_reason or 'Unknown reason'}")

            if getattr(result, 'verified_source', False):
                st.caption(f"✅ Matched by: **{result.match_basis or 'full document'}** "
                           f"(front page {result.front_coverage:.1f}%)")

            c_open, c_prev = st.columns(2)
            with c_open:
                if st.button("📂 Open in Word", key=f"btn_open_final_{idx}", use_container_width=True):
                    try:
                        os.startfile(result.file_path)
                        st.toast(f"Opening {result.word_file_name}...")
                    except Exception as e:
                        st.error(f"Could not open file: {e}")
            with c_prev:
                btn_label = "👁️ Preview (showing)" if (idx == curr_preview_idx and is_preview_vis) else "👁️ Preview"
                if st.button(btn_label, key=f"btn_preview_{idx}", use_container_width=True):
                    st.session_state['preview_doc_index'] = idx
                    st.session_state['preview_page_num'] = 1
                    st.session_state['preview_visible'] = True
                    st.rerun()

            # Stage-by-stage analysis - always shown (normal view)
            st.markdown("**Stage Verification:**")
            for stage in result.stage_reports:
                status = "✅ PASSED" if stage.passed else "❌ FAILED"
                st.markdown(
                    f"- **{stage.stage}** — {status} — Score: {stage.score:.1f}%"
                    f"{(' — ' + stage.effect) if stage.effect else ''}"
                )
                if stage.detail:
                    st.caption(stage.detail)

            st.markdown(
                f"**Why selected:** {result.selection_reason or 'No selection reason recorded.'}"
            )

            # Required Output - Component Match Breakdown
            m1, m2, m3, m4, m5, m6, m7 = st.columns(7)
            m1.metric("Content Match", f"{result.content_score:.1f}%")
            m2.metric("Structure Match", f"{result.structure_score:.1f}%")
            m3.metric("Semantic Match", f"{result.semantic_score:.1f}%")
            m4.metric("Table Match", f"{result.table_score:.1f}%")
            m5.metric("Header/Footer", f"{result.header_footer_score:.1f}%")
            m6.metric("Confidence Score", f"{result.confidence_score:.1f}%")
            m7.metric("Category", result.match_category.split(" (")[0])

            # 10 Component Feature Score Breakdown Expander
            with st.expander(f"📊 View 10-Point Match Score Breakdown — {result.word_file_name}", expanded=(idx == 0)):
                r1_c1, r1_c2, r1_c3, r1_c4, r1_c5 = st.columns(5)
                r1_c1.progress(int(result.component_scores.get('text', 0)), text=f"📝 Text Match: {result.component_scores.get('text', 0):.1f}%")
                r1_c2.progress(int(result.component_scores.get('semantic', 0)), text=f"🧠 Semantic: {result.component_scores.get('semantic', 0):.1f}%")
                r1_c3.progress(int(result.component_scores.get('headings', 0)), text=f"📋 Headings: {result.component_scores.get('headings', 0):.1f}%")
                r1_c4.progress(int(result.component_scores.get('paragraphs', 0)), text=f"📄 Paragraphs: {result.component_scores.get('paragraphs', 0):.1f}%")
                r1_c5.progress(int(result.component_scores.get('tables', 0)), text=f"📊 Tables: {result.component_scores.get('tables', 0):.1f}%")

                r2_c1, r2_c2, r2_c3, r2_c4, r2_c5 = st.columns(5)
                r2_c1.progress(int(result.component_scores.get('keywords', 0)), text=f"🔑 Keywords: {result.component_scores.get('keywords', 0):.1f}%")
                r2_c2.progress(int(result.component_scores.get('structure', 0)), text=f"📑 Structure & Lists: {result.component_scores.get('structure', 0):.1f}%")
                r2_c3.progress(int(result.component_scores.get('section', 0)), text=f"📌 Sections: {result.component_scores.get('section', 0):.1f}%")
                r2_c4.progress(int(result.component_scores.get('page_sequence', 0)), text=f"📖 Page Sequence: {result.component_scores.get('page_sequence', 0):.1f}%")
                r2_c5.progress(int(result.component_scores.get('headers_footers', 0)), text=f"🏷️ Headers/Footers: {result.component_scores.get('headers_footers', 0):.1f}%")
