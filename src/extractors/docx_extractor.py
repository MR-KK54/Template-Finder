import os
import tempfile
import uuid
import zipfile
import re
import docx
from typing import List, Optional, Set
from src.models import DocumentFeatures, TableData, PageData
from src.indexer.fingerprint import build_fingerprint
from src.utils.logger import get_logger
from src.utils.file_utils import canonical_path, compute_file_hash, to_long_path

logger = get_logger("docx_extractor")

def convert_doc_to_docx(filepath: str) -> Optional[str]:
    """Converts a legacy binary .doc file to .docx using Microsoft Word COM
    automation. Returns the path of the converted file, or None if Word is
    unavailable or conversion fails."""
    try:
        import win32com.client
    except Exception:
        logger.warning("pywin32 not installed - cannot convert .doc files.")
        return None

    converted = None
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        tmp_dir = tempfile.mkdtemp(prefix="doc_convert_")
        converted = os.path.join(tmp_dir, f"{uuid.uuid4().hex}.docx")
        doc.SaveAs2(converted, FileFormat=12)  # wdFormatXMLDocument (.docx)
        doc.Close(False)
        logger.info(f"Converted legacy .doc -> .docx via Word COM: {filepath}")
    except Exception as e:
        logger.error(f"Word COM conversion failed for {filepath}: {e}")
        if converted and os.path.exists(converted):
            try:
                os.remove(converted)
            except Exception:
                pass
        converted = None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
    return converted

def extract_doc_text_via_com(filepath: str) -> str:
    """Last-resort plain-text extraction of a .doc file via Word COM."""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
            try:
                return doc.Content.Text or ""
            finally:
                doc.Close(False)
        finally:
            word.Quit()
    except Exception as e:
        logger.error(f"Word COM text extraction failed for {filepath}: {e}")
        return ""

class DocxExtractor:
    """Extractor for Microsoft Word (.docx) documents."""

    def extract(self, filepath: str) -> DocumentFeatures:
        long_filepath = to_long_path(filepath)
        target_path = long_filepath if os.path.exists(long_filepath) else filepath

        if not os.path.exists(target_path):
            raise FileNotFoundError(f"Word document file not found: {filepath}")

        canonical = canonical_path(filepath)
        filename = os.path.basename(canonical)
        folder_name = os.path.basename(os.path.dirname(canonical))

        try:
            file_size = os.path.getsize(target_path)
            last_modified = os.path.getmtime(target_path)
        except Exception:
            file_size = 0
            last_modified = 0.0

        file_hash = compute_file_hash(filepath)

        headings: List[str] = []
        paragraphs: List[str] = []
        lists: List[str] = []
        tables: List[TableData] = []
        keywords: Set[str] = set()
        full_text_chunks: List[str] = []

        # Legacy binary .doc files need conversion before python-docx can read them
        parse_path = target_path
        if canonical.lower().endswith(".doc"):
            converted_path = convert_doc_to_docx(canonical)
            if converted_path:
                parse_path = converted_path
            else:
                plain_text = extract_doc_text_via_com(canonical)
                if plain_text.strip():
                    paragraphs = [p.strip() for p in plain_text.splitlines() if p.strip()]
                    full_text_chunks = list(paragraphs)
                    for line in paragraphs:
                        words = [w.lower().strip(".,!?;:()[]\"'") for w in line.split()]
                        keywords.update([w for w in words if len(w) >= 4 and w.isalpha()])
                    full_text = "\n\n".join(full_text_chunks)
                    section_count = max(1, len(headings) + 1)
                    page_data = PageData(page_num=1, text=full_text, headings=headings, tables=tables, is_scanned=False)
                    doc_features = DocumentFeatures(
                        filepath=canonical, filename=filename, folder_name=folder_name,
                        file_size=file_size, last_modified=last_modified, file_hash=file_hash,
                        full_text=full_text, headings=headings, paragraphs=paragraphs,
                        tables=tables, lists=lists, headers_footers=[], keywords=keywords,
                        page_count=1, section_count=section_count, pages=[page_data],
                        is_scanned_pdf=False
                    )
                    doc_features.fingerprint = build_fingerprint(doc_features)
                    return doc_features
                else:
                    logger.error(f"Legacy .doc file could not be converted or read: {filepath}")
                    full_text_chunks = [f"[Unreadable Document: {filename}]"]
                    headers_footers = []
                    page_count = 1

        try:
            doc = docx.Document(parse_path)

            # 1. Paragraphs, Headings, Lists
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue

                full_text_chunks.append(text)
                style_name = (p.style.name or "").lower()

                if "heading" in style_name or "title" in style_name or "subtitle" in style_name:
                    headings.append(text)
                elif "list" in style_name or text.startswith(("- ", "* ", "• ")) or (len(text) > 2 and text[0].isdigit() and text[1] in ".):"):
                    lists.append(text)
                else:
                    paragraphs.append(text)

                # Collect keywords
                words = [w.lower().strip(".,!?;:()[]\"'") for w in text.split()]
                keywords.update([w for w in words if len(w) >= 4 and w.isalpha()])

            # 2. Tables
            for t in doc.tables:
                rows_cnt = len(t.rows)
                cols_cnt = len(t.columns) if rows_cnt > 0 else 0
                cell_text: List[List[str]] = []
                headers: List[str] = []

                for row_idx, row in enumerate(t.rows):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    cell_text.append(row_cells)
                    if row_idx == 0:
                        headers = row_cells
                    for c_text in row_cells:
                        if c_text:
                            full_text_chunks.append(c_text)

                flat_text = " ".join([c for r in cell_text for c in r if c])
                tables.append(TableData(
                    rows=rows_cnt,
                    cols=cols_cnt,
                    headers=headers,
                    cell_text=cell_text,
                    flat_text=flat_text
                ))

            # 3. Headers and Footers
            header_parts: List[str] = []
            footer_parts: List[str] = []
            for section in doc.sections:
                try:
                    if section.header:
                        for hp in section.header.paragraphs:
                            if hp.text.strip():
                                header_parts.append(hp.text.strip())
                    if section.footer:
                        for fp in section.footer.paragraphs:
                            if fp.text.strip():
                                footer_parts.append(fp.text.strip())
                except Exception:
                    pass

            # 3b. Fallback: read header/footer parts directly from the docx
            # package. python-docx sometimes fails to expose header parts via
            # the section API (e.g. address letterheads), even though the parts
            # exist in the package and are rendered on every page.
            if not header_parts and not footer_parts:
                header_parts, footer_parts = self._extract_headers_footers_from_package(parse_path)

            headers_footers = header_parts + footer_parts

            # Page count estimation for Word docs: len(doc.sections) is
            # unreliable (most documents use a single section regardless of
            # length), so estimate from content volume instead.
            word_count = len(" ".join(full_text_chunks).split())
            page_count = max(1, round(word_count / 300.0)) if word_count else 1
        except Exception as e:
            logger.error(f"Error parsing Word file {filepath}: {e}")
            full_text_chunks = [f"[Unreadable Document: {filename}]"]
            headers_footers = []
            header_parts = []
            footer_parts = []
            page_count = 1

        # Headers render at the TOP of every exported PDF page and footers at
        # the BOTTOM - mirror that order so ordered coverage can match the
        # PDF text against the Word document.
        full_text = "\n\n".join(header_parts + full_text_chunks + footer_parts)

        # Estimated section count (consistent with PDF extractor): one section
        # per heading-like block plus the document body.
        section_count = max(1, len(headings) + 1)

        # Create single synthesized page data for uniform feature comparison
        page_data = PageData(
            page_num=1,
            text=full_text,
            headings=headings,
            tables=tables,
            is_scanned=False
        )

        doc_features = DocumentFeatures(
            filepath=canonical,
            filename=filename,
            folder_name=folder_name,
            file_size=file_size,
            last_modified=last_modified,
            file_hash=file_hash,
            full_text=full_text,
            headings=headings,
            paragraphs=paragraphs,
            tables=tables,
            lists=lists,
            headers_footers=headers_footers,
            keywords=keywords,
            page_count=page_count,
            section_count=section_count,
            pages=[page_data],
            is_scanned_pdf=False
        )
        doc_features.fingerprint = build_fingerprint(doc_features)
        return doc_features

    def _extract_headers_footers_from_package(self, docx_path: str):
        """Reads word/header*.xml and word/footer*.xml parts directly from the
        docx package and returns (headers, footers) visible text lists
        (deduplicated). Falls back to this when python-docx fails to expose
        header/footer parts."""
        headers: List[str] = []
        footers: List[str] = []
        try:
            with zipfile.ZipFile(docx_path) as zf:
                for n in zf.namelist():
                    if not re.match(r"word/(header|footer)\d*\.xml$", n):
                        continue
                    xml = zf.read(n).decode("utf-8", "ignore")
                    chunk = " ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)).strip()
                    if not chunk:
                        continue
                    if n.startswith("word/header"):
                        if chunk not in headers:
                            headers.append(chunk)
                    else:
                        if chunk not in footers:
                            footers.append(chunk)
        except Exception as e:
            logger.warning(f"Package header/footer extraction failed for {docx_path}: {e}")
        return headers, footers
